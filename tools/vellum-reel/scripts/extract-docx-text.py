from pathlib import Path
import sys

from docx import Document


def extract(path: Path) -> str:
    doc = Document(path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        lines.append(f"[TABLE {table_index}]")
        for row in table.rows:
            lines.append("\t".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))
    return "\n".join(lines) + "\n"


for raw in sys.argv[1:]:
    source = Path(raw)
    destination = Path(".cache/docx-review") / f"{source.stem}.txt"
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(extract(source), encoding="utf-8")
    print(destination)
